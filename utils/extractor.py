# utils/extractor.py
# Handles text extraction from uploaded PDF and DOCX resume files.
# Uses pdfplumber for PDFs and python-docx for DOCX files.

import io
import re


def extract_text_from_pdf(file_bytes):
    """
    Extract all text from a PDF file using pdfplumber.

    pdfplumber is preferred over PyPDF2 for better whitespace handling
    and more accurate layout-aware extraction.

    Args:
        file_bytes (bytes): Raw bytes of the PDF file.

    Returns:
        str: Full extracted text joined across all pages.

    Raises:
        ValueError: If extraction fails.
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is not installed. Run: pip install pdfplumber")

    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to read PDF (page error): {str(e)}")

    if not text.strip():
        raise ValueError(
            "No readable text found in this PDF. "
            "The file may be scanned or image-based. "
            "Please use a text-based PDF or DOCX instead."
        )

    return text.strip()


def extract_text_from_docx(file_bytes):
    """
    Extract all text from a DOCX file using python-docx.

    Extracts text from both paragraphs and tables to capture
    all structured content in the document.

    Args:
        file_bytes (bytes): Raw bytes of the DOCX file.

    Returns:
        str: Full extracted text.

    Raises:
        ValueError: If extraction fails.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))

        # Extract from paragraphs (main body, headings, bullet points)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text.strip() + "\n"

        # Extract from tables (some resumes store skills/experience in tables)
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    text += "  ".join(row_texts) + "\n"

    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {str(e)}")

    if not text.strip():
        raise ValueError("No readable text found in the DOCX file.")

    return text.strip()


def extract_text(file_bytes, file_extension):
    """
    Dispatcher: extract text from a file based on its extension.

    Args:
        file_bytes (bytes): Raw bytes of the uploaded file.
        file_extension (str): File extension, e.g., 'pdf' or 'docx'.

    Returns:
        str: Extracted plain text content.

    Raises:
        ValueError: If the file format is unsupported.
    """
    ext = file_extension.lower().strip(".")

    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file format: '{ext}'. Please upload a PDF or DOCX file."
        )


def get_resume_sections(text):
    """
    Detect common resume sections by scanning for known section header keywords.

    This is used by the scorer to determine which sections are present
    and award points accordingly.

    Args:
        text (str): Raw resume text.

    Returns:
        dict: Mapping of section_name -> True if detected.
              Example: {"education": True, "experience": True, ...}
    """
    section_keywords = {
        "education": [
            "education", "academic background", "qualification",
            "degree", "university", "college", "school", "academics"
        ],
        "experience": [
            "experience", "work history", "employment", "work experience",
            "professional experience", "internship", "intern", "career"
        ],
        "skills": [
            "skills", "technical skills", "core competencies",
            "technologies", "expertise", "proficiencies"
        ],
        "projects": [
            "projects", "personal projects", "academic projects",
            "portfolio", "side projects", "key projects"
        ],
        "certifications": [
            "certification", "certificate", "courses", "training",
            "credential", "licensed", "certified"
        ],
        "summary": [
            "summary", "objective", "profile", "about me",
            "career objective", "professional summary", "overview"
        ],
        "contact": [
            "email", "phone", "mobile", "linkedin", "github",
            "contact", "address", "portfolio"
        ],
        "achievements": [
            "achievement", "award", "honor", "recognition",
            "accomplishment", "distinction"
        ]
    }

    text_lower = text.lower()
    detected = {}

    for section, keywords in section_keywords.items():
        for keyword in keywords:
            if keyword in text_lower:
                detected[section] = True
                break

    return detected